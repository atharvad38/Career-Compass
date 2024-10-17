
# import os
# import fitz  # PyMuPDF for PDF processing
# import streamlit as st
# from groq import Groq
# from docx import Document  # Library to handle DOCX files
# import nltk
# from io import BytesIO

# # Download NLTK resources (run once)
# nltk.download('punkt')

# # Get the API key from environment variable
# groq_api_key = os.getenv("GROQ_API_KEY")

# # Check if the API key was retrieved successfully
# if not groq_api_key:
#     st.error("API key not found. Please set the GROQ_API_KEY environment variable in Streamlit Cloud.")

# # Streamlit application title and description
# st.title("Career Compass")
# st.write("### Upload your resume and specify the job title to get a professional score tailored to your target role.")

# # Streamlit input field for job title
# job_title = st.text_input("Enter the job title for which you want to rate your resume (e.g., Data Scientist, Software Engineer, Marketing Head, etc.):")

# # Streamlit file uploader for PDF or DOCX files
# uploaded_file = st.file_uploader("Upload your resume (PDF or DOCX)", type=["pdf", "docx"])

# # Function to extract text from PDF
# def extract_text_from_pdf(uploaded_file):
#     """Extract text from a PDF file using PyMuPDF (fitz)."""
#     pdf_data = BytesIO(uploaded_file.read())
#     doc = fitz.open(stream=pdf_data, filetype="pdf")
#     text = ""
#     for page_num in range(doc.page_count):
#         page = doc.load_page(page_num)
#         text += page.get_text()
#     doc.close()
#     return text

# # Function to extract text from DOCX
# def extract_text_from_docx(uploaded_file):
#     """Extract text from a DOCX file using python-docx."""
#     doc = Document(uploaded_file)
#     text = []
#     for para in doc.paragraphs:
#         text.append(para.text)
#     return '\n'.join(text)

# # Function to score the resume with Groq
# def score_resume_with_groq(text, job_title):
#     """Send resume data to Groq API for scoring."""
#     client = Groq(api_key=groq_api_key)

#     prompt = f"Give me a resume score for this resume with a focus on an ideal candidate for the role of {job_title}. Provide a score out of 10 for the overall quality of the resume and a one-line description and also give recommended job title:Don't hesitate to give less score if job title ({job_title}) and the resume are not at all related else if they are related give good score, also check for education properly. Format: Resume score: x/10\nDescription: (1 line):\n\n{text}\n\nRecommended Job Title:"

#     chat_completion = client.chat.completions.create(
#         messages=[
#             {
#                 "role": "user",
#                 "content": prompt,
#             }
#         ],
#         model="llama3-8b-8192",
#     )

#     return chat_completion.choices[0].message.content.strip()

# if uploaded_file and job_title:
#     # Determine file type and extract text accordingly
#     if uploaded_file.type == "application/pdf":
#         resume_text = extract_text_from_pdf(uploaded_file)
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         resume_text = extract_text_from_docx(uploaded_file)
#     else:
#         st.error("Unsupported file type. Please upload a PDF or DOCX file.")

#     # Check if resume text was extracted successfully
#     if resume_text:
#         st.write("Processing your resume for the role of", job_title, "... Please wait.")
#         # Get the resume score using Groq API
#         score = score_resume_with_groq(resume_text, job_title)
#         st.write(f"**Resume Score:** {score}")
#     else:
#         st.error("Failed to extract text from the uploaded resume.")
# else:
#     st.info("Please upload a resume and enter the job title.")

import os
import fitz  # PyMuPDF for PDF processing
import streamlit as st
from groq import Groq, AuthenticationError
from docx import Document  # Library to handle DOCX files
import nltk
from io import BytesIO

# Download NLTK resources (run once)
nltk.download('punkt')

# Get the API key from environment variable
groq_api_key = os.getenv("GROQ_API_KEY")

# Check if the API key was retrieved successfully
if not groq_api_key:
    st.error("API key not found. Please set the GROQ_API_KEY environment variable in Streamlit Cloud.")

# Streamlit application title and description
st.title("Career Compass")
st.write("### Upload your resume and specify the job title to get a professional score tailored to your target role.")

# Streamlit input field for job title
job_title = st.text_input("Enter the job title for which you want to rate your resume (e.g., Data Scientist, Software Engineer, Marketing Head, etc.):")

# Streamlit file uploader for PDF or DOCX files
uploaded_file = st.file_uploader("Upload your resume (PDF or DOCX)", type=["pdf", "docx"])

# Function to extract text from PDF
def extract_text_from_pdf(uploaded_file):
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    pdf_data = BytesIO(uploaded_file.read())
    doc = fitz.open(stream=pdf_data, filetype="pdf")
    text = ""
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text += page.get_text()
    doc.close()
    return text

# Function to extract text from DOCX
def extract_text_from_docx(uploaded_file):
    """Extract text from a DOCX file using python-docx."""
    doc = Document(uploaded_file)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return '\n'.join(text)

# Function to score the resume with Groq
def score_resume_with_groq(text, job_title):
    """Send resume data to Groq API for scoring."""
    client = Groq(api_key=groq_api_key)

    prompt = f"Give me a resume score for this resume with a focus on an ideal candidate for the role of {job_title}. Provide a score out of 10 for the overall quality of the resume and a one-line description and also give recommended job title: Don't hesitate to give less score if job title ({job_title}) and the resume are not at all related else if they are related give good score, also check for education properly. Format: Resume score: x/10\nDescription: (1 line):\n\n{text}\n\nRecommended Job Title:"

    try:
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-8b-8192",
        )
        return chat_completion.choices[0].message.content.strip()
    except AuthenticationError:
        st.error("Authentication failed. Please check your Groq API key.")
        return None
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
        return None

# Main logic to process resume and score it
if uploaded_file and job_title:
    # Determine file type and extract text accordingly
    if uploaded_file.type == "application/pdf":
        resume_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        resume_text = extract_text_from_docx(uploaded_file)
    else:
        st.error("Unsupported file type. Please upload a PDF or DOCX file.")

    # Check if resume text was extracted successfully
    if resume_text:
        st.write(f"Processing your resume for the role of {job_title}... Please wait.")
        # Get the resume score using Groq API
        score = score_resume_with_groq(resume_text, job_title)
        if score:
            st.write(f"**Resume Score:** {score}")
    else:
        st.error("Failed to extract text from the uploaded resume.")
else:
    st.info("Please upload a resume and enter the job title.")
